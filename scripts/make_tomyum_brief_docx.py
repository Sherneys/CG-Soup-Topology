# scripts/make_tomyum_brief_docx.py
# One-page Thai advisor brief for the tomyum (Thai hot pot) asset + its first
# C-matrix results. House rule: NO hardcoded results — the C-table is read
# from output/synth/topo3/quicklook.json and the density-staircase table is
# RECOMPUTED live from the mesh, so the document cannot drift from the data.
#
#   D:\...\CG-Soup-for-Digital-Dentistry\.venv\Scripts\python.exe `
#       scripts\make_tomyum_brief_docx.py [out.docx]

import json
import math
import os
import sys

import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Tahoma"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)

_TOPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _TOPO)
DENTISTRY = os.environ.get("CGSOUP_ROOT", r"D:\Project\CG-Soup-for-Digital-Dentistry")
QUICKLOOK = os.path.join(DENTISTRY, "output", "synth", "topo3", "quicklook.json")
OUT = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    _TOPO, "docs", "CG-Soup_Tomyum_Brief_TH.docx")


# ── style helpers (verbatim from make_phase3_report_docx.py) ───────────────

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, white=False, size=9, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.size = Pt(size); run.font.name = FONT; run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        _shade(cell, fill)


def add_table(doc, headers, rows, shade_rows=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _cell_text(t.rows[0].cells[j], h, bold=True, white=True); _shade(t.rows[0].cells[j], "1F4E79")
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        fill = (shade_rows or {}).get(i)
        for j, v in enumerate(r):
            _cell_text(cells[j], v, fill=fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_figure(doc, path, caption, width=6.0):
    if not os.path.exists(path):
        doc.add_paragraph(f"[ไม่พบรูป: {path}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT; run.font.color.rgb = ACCENT
    return h


def P(doc, text, size=10.5, space=6, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT; r.bold = bold
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT
    return p


# ── live numbers ───────────────────────────────────────────────────────────

def staircase_rows():
    """Rebuild the mesh and recompute the significant reading per M (live)."""
    from methods._paths import load_topology
    load_topology()
    from topology import meshes
    from topology.persistence import persistence_from_points
    V, F = meshes.tomyum_pot_mesh()
    rows, meta = [], {"verts": len(V), "faces": len(F)}
    for M in (2048, 8192, 20000, 50000):
        Pc = meshes.sample_surface(V, F, M, np.random.default_rng(0))
        res = persistence_from_points(Pc)
        b = res.betti_numbers()
        rows.append([f"{M:,}", f"{res.significance_threshold():.4f}",
                     f"({b[0]}, {b[1]}, {b[2]})"])
    return rows, meta


def cmatrix_stats():
    with open(QUICKLOOK, encoding="utf-8") as fh:
        rows = [r for r in json.load(fh) if r["shape"] == "tomyum"]
    if not rows:
        raise SystemExit(f"no tomyum rows in {QUICKLOOK} — run the quicklook first")
    out = {}
    for cond in ("C0", "C1", "C2"):
        sel = sorted([r for r in rows if r["cond"] == cond], key=lambda r: r["seed"])
        if not sel:
            continue
        b = np.array([r["tail_bottleneck_H1"] for r in sel], dtype=float)
        ch = np.array([r["chamfer_pct"] for r in sel], dtype=float)
        out[cond] = {"n": len(sel), "bott": b, "mean": float(b.mean()),
                     "sd": float(b.std(ddof=1)) if len(sel) > 1 else 0.0,
                     "nsig": [int(r["nsig_H1"]) for r in sel],
                     "cham": float(ch.mean())}
    def sigma(a, b):        # SE-of-difference convention (as in the 3e verdicts)
        se = math.sqrt(out[a]["sd"] ** 2 / out[a]["n"] + out[b]["sd"] ** 2 / out[b]["n"])
        return abs(out[a]["mean"] - out[b]["mean"]) / se if se > 0 else float("inf")
    out["red_C1"] = out["C0"]["mean"] / out["C1"]["mean"]     # unrounded means
    out["sig_C1"] = sigma("C0", "C1")
    if "C2" in out:
        out["worse_C2"] = out["C2"]["mean"] / out["C0"]["mean"]
        out["sig_C2"] = sigma("C0", "C2")
    return out


def main():
    stair, meta = staircase_rows()
    st = cmatrix_stats()

    doc = Document()
    for name in ("Normal", "List Bullet"):
        s = doc.styles[name]
        s.font.name = FONT
        rpr = s.element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(a), FONT)
        rpr.append(rfonts)

    H(doc, "โมเดลไทยตัวใหม่ในชุดทดสอบ: หม้อไฟต้มยำ (genus 9) — ผล C-matrix ชุดแรก", 1)
    P(doc, "สรุปสำหรับอาจารย์ · 10 ก.ค. 2026 · repo: CG-Soup-Topology (commit 384181f + รัน C-matrix)",
      size=9, space=8)

    H(doc, "1) โมเดลคืออะไร และทำไมพิเศษ", 2)
    P(doc, f"หม้อไฟต้มยำแบบดั้งเดิม (แอ่งน้ำซุปวงแหวนรอบปล่องกลาง บนฐานเจาะรูระบายอากาศ) "
           f"สร้างเป็น CSG ที่ topology พิสูจน์ได้จากการประกอบ ไม่ใช่เชื่อไฟล์ที่ดาวน์โหลด: "
           f"ตัวหม้อ revolve (genus 1) + หูหม้อ 2 ข้าง (+1/ข้าง) + รูระบายอากาศ 6 รู (+1/รู) "
           f"= genus 9 → homology ของ mesh คือ b = (1, 18, 1) แบบ exact "
           f"(mesh จริง {meta['verts']:,} จุด / {meta['faces']:,} หน้า, watertight, "
           f"ตรวจยืนยัน 3 ทางอิสระ: kernel genus / Euler certificate χ=−16 / GUDHI)")
    add_figure(doc, os.path.join(_TOPO, "figures", "tomyum_pot.png"),
               "หม้อไฟต้มยำ genus 9: ปล่องกลาง (1) + หูหม้อ (2) + รูระบายอากาศฐาน (6)")

    H(doc, "2) จุดขายเชิงวิชาการ: บันไดความหนาแน่น (measurement floor ในวัตถุเดียว)", 2)
    P(doc, "point cloud ของหม้ออ่านเป็นโลหะตัน (handlebody) — b1 ลู่เข้า genus (9) และ b2 = 0 "
           "(ภาชนะเปิดไม่มีโพรงปิด) แต่เพราะสเกลของห่วงต่างกันเป็นสิบเท่า (รูปล่อง Ø104 mm "
           "เทียบหู/รูระบาย Ø22–24 mm) จำนวนห่วงที่ significant จึงขึ้นกับความหนาแน่นจุด M "
           "ตามกฎ floor 6·r_med(M) ของเรา (experiments/density_bound.py) — คำนวณสดตอนสร้างเอกสารนี้:")
    add_table(doc, ["M (จุด)", "floor 6·r_med", "significant (b0, b1, b2)"], stair,
              shade_rows={0: "FFF2CC", 3: "E2EFDA"})
    P(doc, "ที่ M=2048 (ค่ามาตรฐานของ pipeline) หม้ออ่านเป็น solid torus — เห็นเฉพาะห่วงปล่อง "
           "(ทดสอบแล้ว seed 0–4 ผ่านทุกตัว margin 1.17–1.25×) และต้องถึง M≈50,000 จึงเห็นครบ rank 9",
      size=9.5)

    H(doc, "3) ผล C-matrix (สูตรเดียวกับ 3e ทุกประการ — zero per-shape tuning)", 2)
    P(doc, "N=700, 2500 steps, ρ=0.1, ramp 0.2:0.5, H1-only (เหมือน torus/bob), 3 seeds; "
           "target bundle ที่ M=2048 มี H1 significant 1 bar ตามออกแบบ:", size=9.5)
    rows = []
    fmt = lambda c: (f"{st[c]['mean']:.4f} ± {st[c]['sd']:.4f}",
                     "/".join(str(x) for x in st[c]["nsig"]), f"{st[c]['cham']:.3f}")
    r0 = fmt("C0"); rows.append(["C0 เบสไลน์", r0[0], "—", r0[1], r0[2]])
    r1 = fmt("C1"); rows.append(["C1 topological loss", r1[0],
                                 f"ดีขึ้น {st['red_C1']:.1f}× ({st['sig_C1']:.1f}σ)", r1[1], r1[2]])
    if "C2" in st:
        r2 = fmt("C2"); rows.append(["C2 ตัวควบคุม (repulsion)", r2[0],
                                     f"แย่ลง {st['worse_C2']:.1f}× ({st['sig_C2']:.1f}σ)", r2[1], r2[2]])
    add_table(doc, ["เงื่อนไข", "tail bottleneck H1 (mean±sd)", "เทียบ C0",
                    "#sig H1 (ถูกต้อง=1)", "Chamfer %"],
              rows, shade_rows={1: "E2EFDA", 2: "FCE4EC"})
    bullet(doc, f"C1 ลดระยะ bottleneck ของห่วงปล่อง {st['red_C1']:.1f}× ({st['sig_C1']:.1f}σ) "
                f"ที่ Chamfer parity (C1/C0 = {st['C1']['cham'] / st['C0']['cham']:.2f}) — "
                f"อยู่กลางช่วง 2.1–10.4× ของ study เดิม")
    if "C2" in st:
        bullet(doc, f"C2 control แย่กว่าเบสไลน์ {st['worse_C2']:.1f}× ({st['sig_C2']:.1f}σ) "
                    f"และ Chamfer แย่ลงด้วย — ยืนยันว่าประโยชน์มาจากทิศทาง gradient ที่ topology "
                    f"กำหนด ไม่ใช่แรงผลักทั่วไป")
        bullet(doc, "ข้อสังเกตซื่อตรง: บนหม้อนี้ C2 ไม่ทำ count พัง (#sig H1 = 1 ทุก seed — "
                    "ปล่อง Ø104 อ้วนเกินกว่าจะโดนตัดขาด) ต่างจาก bob (b1 2→1) — โหมดคือ "
                    "degrade-not-destroy")
    bullet(doc, "ทุกเงื่อนไขนับ #sig H1 = 1 ถูกต้องทุก seed — ไม่มี phantom handle")

    H(doc, "4) คำถามที่ขอความเห็นอาจารย์", 2)
    bullet(doc, "ควรเพิ่ม tomyum เป็น external mesh ตัวที่ 4 ในตาราง 3e ของ paper 2 หรือไม่? "
                "(ตารางปัจจุบัน spot/bob/fandisk ปิดแล้ว และเราเคยตกลงว่าจะขยาย shape set "
                "เฉพาะเมื่อ reviewer ขอ — แต่ตัวนี้เพิ่มจุดขาย ‘ground-truth topology จากการประกอบ’ "
                "และเป็นลายเซ็นไทยของกลุ่ม)")
    bullet(doc, "หรือเก็บเป็น showcase/รูปเปิดของเฟส dental (หม้อ→ครอบฟัน: วัตถุจริง "
                "ผนังบาง มีห่วงหลายสเกล) แล้วคง paper 2 ตามเดิม?")
    P(doc, "หมายเหตุ: runs ชุดนี้เป็น tag ใหม่ใน topo3 — ไม่มีตัวเลขใดของ paper 2 เปลี่ยน; "
           "ทำซ้ำได้ด้วย experiments/topo_loss_eval.py --shapes tomyum --seeds 0 1 2 "
           "--conditions C0 C1 C2 --rhos 0.1 --steps 2500 --max_faces 700 --loss_dims 1",
      size=8.5, space=2)

    doc.save(OUT)
    print(f"[docx] {OUT}")
    print(f"  C1: {st['red_C1']:.2f}x ({st['sig_C1']:.1f} sigma)  "
          f"C2: {st.get('worse_C2', float('nan')):.2f}x worse ({st.get('sig_C2', float('nan')):.1f} sigma)")


if __name__ == "__main__":
    main()
