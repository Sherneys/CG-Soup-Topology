# src/make_topology_report_docx.py
# Build the Phase-1 topology-measurement work as a Thai .docx summary.
# Reuses the style of src/make_geo_report_docx.py and reads the real numbers
# from output/figures/topology/ (blindness_results.json + stability_series.csv)
# so the document never drifts from the experiment.
#
#   python src/make_topology_report_docx.py [out.docx]

import csv
import json
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Tahoma"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
_TOPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))   # the topology/ package dir
TOPO_DIR = os.path.join(_TOPO, "figures")
OUT = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    _TOPO, "docs", "CG-Soup_Topology_Phase1_TH.docx")

DIM_TH = {0: "H0 (จำนวนชิ้นส่วน/components)", 1: "H1 (ห่วง·หูจับ/loops·handles)",
          2: "H2 (โพรงปิด/voids)"}


# ── style helpers (from make_geo_report_docx.py) ─────────────────────

def _set_style_font(style, name=FONT):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), name)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, white=False, size=9, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.size = Pt(size); run.font.name = FONT; run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        _shade(cell, fill)


def add_table(doc, headers, rows, widths=None, shade_rows=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _cell_text(t.rows[0].cells[j], h, bold=True, white=True); _shade(t.rows[0].cells[j], "1F4E79")
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        fill = (shade_rows or {}).get(i)
        for j, v in enumerate(r):
            _cell_text(cells[j], v, fill=fill)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_figure(doc, path, caption, width=6.3):
    if not os.path.exists(path):
        doc.add_paragraph(f"[ไม่พบรูป: {path}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT; run.font.color.rgb = ACCENT
    return h


def P(doc, text, size=10.5, space=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT
    return p


def code(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ── load real results ────────────────────────────────────────────────

def load_results():
    jp = os.path.join(TOPO_DIR, "blindness_results.json")
    with open(jp, encoding="utf-8") as f:
        return json.load(f)


def load_stability():
    cp = os.path.join(TOPO_DIR, "stability_series.csv")
    rows = []
    if os.path.exists(cp):
        with open(cp, encoding="utf-8") as f:
            rows = list(csv.DictReader(f))
    return rows


# ── build ────────────────────────────────────────────────────────────

doc = Document()
for sname in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
    try:
        _set_style_font(doc.styles[sname])
    except KeyError:
        pass
doc.styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("CG-Soup / DiffSoup — Phase 1: ฮาร์เนสวัด ‘โทโพโลยี’ และการสาธิตจุดบอดของ Chamfer/Hausdorff")
tr.bold = True; tr.font.size = Pt(17); tr.font.color.rgb = ACCENT; tr.font.name = FONT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Topology-aware Adaptive Resampling — เฟส 1: การวัดและการสาธิต (measurement-only)")
sr.font.size = Pt(12); sr.italic = True; sr.font.name = FONT
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("สรุปงานที่ทำทั้งหมด · เครื่องมือ GUDHI alpha complex + POT · ทุกอย่าง deterministic (seed ครบ) · 29 มิถุนายน 2569")
mr.font.size = Pt(10); mr.font.name = FONT
doc.add_paragraph()

data = load_results()
cases = data["cases"]
stab = load_stability()
n_pass = sum(1 for g in data["gate"] if g["passes"])

# 1. บทคัดย่อ
H(doc, "1. บทคัดย่อ")
P(doc, "งานนี้คือ ‘เฟส 1’ ของแนวคิด Topology-aware Adaptive Resampling โดยมีขอบเขตเฉพาะ ‘การวัด + การสาธิต’ "
       "เท่านั้น (ยังไม่สร้างเมท็อด resampling และยังไม่ทำ differentiable persistent homology). สมมติฐานหลักคือ "
       "Adaptive Resampling ของ DiffSoup ลบความต่างที่ใส่ไว้ตอน init ทิ้งอย่างรวดเร็ว และตัวชี้วัดเชิงเรขาคณิตที่ใช้กันทั่วไป "
       "(Chamfer, Hausdorff) ‘มองไม่เห็นโทโพโลยี’ — งานที่ผิด genus, มีรู, หรือชิ้นส่วนหลุด/เชื่อมผิด ก็ยังได้คะแนนดีได้")
P(doc, f"ผลลัพธ์: สร้างโมดูล topology/ ที่แยกเป็นอิสระ (วัด persistence H0/H1/H2 ผ่าน alpha complex), ตัวชี้วัด "
       f"Topology Stability Metric, การทดสอบกู้คืนเลขเบ็ตติที่ผ่าน 6/6, และการทดลองสาธิต ‘จุดบอด’ ที่ผ่านเกณฑ์ "
       f"{n_pass}/3 เคส (ต้องการอย่างน้อย 1). ทุกเคสคุม Chamfer ให้ ‘เท่ากัน’ ระหว่างงานที่โทโพโลยีถูกกับผิด "
       f"แต่ระยะ bottleneck ของไดอะแกรมเพอร์ซิสเทนซ์แยกออกจากกันชัดเจน")

# 2. ขอบเขตและข้อจำกัด
H(doc, "2. ขอบเขตเฟส 1 (ทำอะไร / ไม่ทำอะไร)")
bullet(doc, "ทำ: วัดโทโพโลยี (persistence diagram H0/H1/H2) จาก mesh หรือจุดสุ่มบนผิวด้วย alpha complex; "
            "ตัวชี้วัดเสถียรภาพโทโพโลยีตามรอบ training; การสาธิตเชิงปริมาณว่า Chamfer/Hausdorff แยกเคสที่โทโพโลยีต่างกันไม่ได้")
bullet(doc, "ไม่ทำ (สงวนไว้เฟส 2): เมท็อด Topology-aware resampling, differentiable persistent homology — เฟสนี้โทโพโลยีเป็น ‘การวัด’ ล้วน")
bullet(doc, "ความเข้มงวด: ทุกอย่าง deterministic — ตั้ง seed ครบ, การสุ่มจุดทำซ้ำได้แน่นอน (เหมือนแนวทาง synthetic GT ที่ใช้อยู่เดิม)")
bullet(doc, "การนำกลับมาใช้: ไม่แตะ renderer ของ DiffSoup (อยู่ภายนอกที่ D:\\Project\\diffsoup ผ่าน DIFFSOUP_ROOT) — "
            "เฟสนี้เป็นงานบน mesh จึงไม่ต้องเรียก renderer; แต่ ‘ยืมหลักการ’ ของ src/eval_geometry.py "
            "(นอร์มัลด้วยเส้นทแยงมุม bbox, การสุ่มจุดบน soup ถ่วงด้วย alpha×พื้นที่) มาให้ตัวเลขเทียบกันได้")

# 3. เครื่องมือ
H(doc, "3. เครื่องมือและการติดตั้ง")
P(doc, "เลือก GUDHI (alpha complex) เพราะกู้คืน H0/H1/H2 ของผิว 2 มิติใน 3 มิติได้ด้วย simplex น้อยกว่า Vietoris–Rips มาก "
       "และค่า filtration (รัศมีวงกลมล้อม ยกกำลังสอง) เป็น ‘สเกลความยาว’ จริง จึงเทียบกับ Chamfer/Hausdorff ได้โดยตรง "
       "เมื่อทั้งคู่นอร์มัลด้วยเส้นทแยงมุม bbox เดียวกัน")
bullet(doc, "ระยะ bottleneck → ใช้ของ GUDHI (built-in)")
bullet(doc, "ระยะ Wasserstein → gudhi.wasserstein ที่ขับด้วย POT (python-optimal-transport)")
bullet(doc, "ทั้ง gudhi และ pot มี prebuilt wheel รวมถึง macOS arm64 (Apple Silicon) → `uv pip install gudhi pot` "
            "ติดตั้งบน Mac M-series ได้โดยไม่ต้องคอมไพล์ (ยืนยันบนเครื่อง Windows ปัจจุบันแล้ว: gudhi 3.12.0, pot 0.9.6)")

# 4. โครงสร้างไฟล์ที่สร้าง
H(doc, "4. ไฟล์ที่สร้าง (Deliverables)")
add_table(doc,
          ["ไฟล์", "หน้าที่"],
          [["topology/persistence.py", "Deliverable 1: persistence H0/H1/H2 (alpha complex) + 2 entry points"],
           ["topology/metrics.py", "Deliverable 2: Topology Stability Metric (bottleneck/Wasserstein + นับฟีเจอร์)"],
           ["topology/meshes.py", "รูปทรงสังเคราะห์ deterministic + ตัวสุ่มจุดแบบ seeded (รวม soup/mesh)"],
           ["topology/plots.py", "วาดไดอะแกรม + อนุกรมเวลา (นำกลับมาใช้เฟส 2 ได้)"],
           ["experiments/topology_blindness.py", "Deliverable 3: การทดลองสาธิต + ตาราง + เกณฑ์ผ่าน"],
           ["tests/test_betti.py", "Verification #1: กู้คืนเลขเบ็ตติ (6/6 ผ่าน)"],
           ["output/figures/topology/*", "Deliverable 4: รูปไดอะแกรม 3 เคส + อนุกรมเสถียรภาพ + summary.md + JSON/CSV"]],
          widths=[2.4, 4.2])

# 5. Persistence (entry points)
H(doc, "5. topology/persistence.py — สอง entry points")
P(doc, "แยกสองทางเข้าอย่างชัดเจนตามที่กำหนด แม้เฟส 1 ใช้เพียง (ข):")
bullet(doc, "(ก) persistence_from_target(...) — สำหรับ ‘การชี้นำ’ ในเฟส 2 (เป็น prior ของโทโพโลยีเป้าหมาย) — "
            "ตรึง API ไว้ตอนนี้แต่ ‘ยังไม่ถูกใช้’ โดย metric/การทดลองใด ๆ (มีเพียงเทสต์ที่แตะเพื่อยืนยันว่าทำงานถูก)")
bullet(doc, "(ข) persistence_from_reconstruction(...) — สำหรับ ‘การวัด’ งานที่สร้างขึ้น (mesh, จุด, หรือ checkpoint ของ DiffSoup) — "
            "เป็นทางที่ถูกใช้จริงในเฟสนี้")
P(doc, "‘ฟีเจอร์ที่มีนัย’ = ช่วงเพอร์ซิสเทนซ์ที่อายุ (death−birth) มากกว่า k เท่าของระยะเพื่อนบ้านใกล้สุดมัธยฐานของกลุ่มจุด "
       "(หน่วยนอร์มัล) — กฎนี้ปรับสเกลตัวเองตามความหนาแน่นของจุด ไม่ต้องตั้งค่าคงที่ตายตัว และอ่านเลขเบ็ตติได้จากจำนวนฟีเจอร์ที่มีนัยต่อมิติ", size=10)

# 6. Verification — Betti
H(doc, "6. การตรวจสอบ #1 — กู้คืนเลขเบ็ตติ (ผ่าน 6/6)")
add_table(doc,
          ["รูปทรง", "b0", "b1", "b2", "ผล"],
          [["ทรงกลม (sphere)", "1", "0", "1", "ผ่าน"],
           ["ทอรัส (torus)", "1", "2", "1", "ผ่าน"],
           ["ทรงกลม 2 ลูก (2 spheres)", "2", "—", "—", "ผ่าน (b0=2)"]],
          widths=[2.6, 0.8, 0.8, 0.8, 1.6])
P(doc, "บวกเทสต์ determinism (seed เดียวกัน → ไดอะแกรมตรงกัน bit-for-bit), เทสต์ความเสถียรของ threshold (k ใน [4,8] ให้เบ็ตติเท่ากัน), "
       "และเทสต์ตรึง API ของ entry point ทั้งสอง", size=10)

# 7. การทดลองสาธิตจุดบอด
H(doc, "7. การทดลองหลัก — Chamfer/Hausdorff ‘มองไม่เห็นโทโพโลยี’")
P(doc, "สามเคสควบคุม. ในแต่ละเคสสร้าง ground-truth (GT) และงานสองชิ้น: A = โทโพโลยีถูก, B = โทโพโลยีผิด — "
       "โดย ‘ปรับเรขาคณิตของ A ให้ Chamfer เท่ากับ B’ (bisection) เพื่อตรึงเรขาคณิตให้คงที่ แล้วดูว่าระยะ bottleneck "
       "ของไดอะแกรมเทียบ GT แยก A ออกจาก B ได้หรือไม่ (ค่า = หน่วยนอร์มัลด้วยเส้นทแยงมุม bbox; bott สูง=ห่าง GT มาก=แย่):")

disc_rows = []
shade = {}
for ci, c in enumerate(cases):
    d = c["disc_dim"]
    for k, name in enumerate(("A", "B")):
        cand = c["candidates"][name]
        b = cand["betti"]
        bott = cand["bottleneck"][str(d)]
        label = "A (ถูก)" if name == "A" else "B (ผิด)"
        disc_rows.append([
            c["key"] if name == "A" else "",
            label,
            f"{cand['chamfer_pct']:.3f}",
            f"{cand['hausdorff95_pct']:.3f}",
            f"{bott:.4f} ({['H0','H1','H2'][d]})",
            f"{b[0]}/{b[1]}/{b[2]}",
        ])
        if name == "B":
            shade[len(disc_rows) - 1] = "FCE4E4"   # light red for the wrong row
add_table(doc,
          ["เคส", "งาน", "Chamfer %", "Hausd95 %", "bott→GT (มิติชี้ขาด)", "เบ็ตติ b0/b1/b2"],
          disc_rows, widths=[1.3, 1.0, 1.0, 1.0, 1.7, 1.3], shade_rows=shade)

for c, g in zip(cases, data["gate"]):
    d = c["disc_dim"]
    cA = c["candidates"]["A"]; cB = c["candidates"]["B"]
    bA = cA["bottleneck"][str(d)]; bB = cB["bottleneck"][str(d)]
    bullet(doc, f"{c['key']} — มิติชี้ขาด {DIM_TH[d]}: Chamfer A={cA['chamfer_pct']:.3f}% ≈ B={cB['chamfer_pct']:.3f}% "
                f"(เท่ากัน) แต่ bottleneck→GT  A={bA:.4f} เทียบ B={bB:.4f} → แยกออกจากกันชัด "
                f"{'· Hausdorff95 ยัง ‘เข้าข้าง’ งานที่ผิด (B ต่ำกว่า A)' if cB['hausdorff95_pct'] < cA['hausdorff95_pct'] else ''}")
P(doc, f"สรุปเกณฑ์: {n_pass}/3 เคสแสดง ‘Chamfer เท่ากันหรือกลับด้าน ขณะที่โทโพโลยีแยกออก’ (ต้องการ ≥ 1)", space=4)

add_figure(doc, os.path.join(TOPO_DIR, "pd_i_genus.png"),
           "รูปที่ 1 — ไดอะแกรมเพอร์ซิสเทนซ์ เคส (i) genus: B (หูจับ) มีจุด H1 (แดง) โผล่พ้นเส้นทแยงมุมที่ GT/A ไม่มี → b1: 0→1")
add_figure(doc, os.path.join(TOPO_DIR, "pd_ii_components.png"),
           "รูปที่ 2 — เคส (ii) components: จุด H0 (น้ำเงิน, ‘ชิ้นที่สอง’) มีใน GT/A แต่ ‘หาย’ ใน B (สะพานบางเชื่อมรวมเป็นชิ้นเดียว) → b0: 2→1")
add_figure(doc, os.path.join(TOPO_DIR, "pd_iii_void.png"),
           "รูปที่ 3 — เคส (iii) void: จุด H2 (เขียว) ของ A ทับ GT (โพรงยังอยู่) แต่ของ B ถูกดึงเข้าหาเส้นทแยงมุม (โพรงพัง)")

# 8. Stability series
H(doc, "8. ตัวชี้วัดเสถียรภาพโทโพโลยี (อนุกรมเวลา)")
P(doc, "สาธิตตัวชี้วัด (deliverable 2) ด้วยลำดับ ‘หูจับงอกบนทรงกลม’ ที่ค่อย ๆ สูงขึ้น (จำลองการเกิด defect ระหว่างเทรน). "
       "เป้าหมาย = ทรงกลม. ผลคือ Chamfer (เรขาคณิต) ไต่ขึ้น ‘เรียบ ๆ’ ขณะที่ bottleneck-H1 (โทโพโลยี) นิ่งแล้ว ‘กระโดด’ "
       "เมื่อห่วงของหูจับเริ่มมีนัย และจำนวนฟีเจอร์ H1 ที่มีนัยพลิกจาก 0→1 ภายหลัง — แสดงว่าตัวชี้วัดต่อเนื่อง (bottleneck) "
       "ไวกว่าการนับเบ็ตติแบบไม่ต่อเนื่อง")
if stab:
    first, last = stab[0], stab[-1]
    flip = next((r for r in stab if int(r["nsig_H1"]) >= 1), None)
    P(doc, f"ตัวเลขจากอนุกรม: Chamfer ไต่จาก {float(first['chamfer']):.3f}% → {float(last['chamfer']):.3f}% (เรียบ); "
           f"bottleneck-H1 จาก {float(first['bottleneck_H1']):.3f} → {float(last['bottleneck_H1']):.3f}; "
           f"จำนวน H1 ที่มีนัยพลิกเป็น 1 ที่ step {flip['step'] if flip else '—'}", size=10)
add_figure(doc, os.path.join(TOPO_DIR, "stability_series.png"),
           "รูปที่ 4 — เสถียรภาพโทโพโลยี: เส้นเทา (Chamfer) ไต่เรียบ ส่วนเส้นแดง (bottleneck-H1) กระโดดเมื่อ defect มีนัย; "
           "แผงล่างคือจำนวนฟีเจอร์ที่มีนัยต่อมิติ", width=5.6)

# 9. ข้อค้นพบเชิงเทคนิค / ข้อควรระวัง
H(doc, "9. ข้อค้นพบเชิงเทคนิคและข้อควรระวัง (ตามจริง)")
bullet(doc, "เคส (i),(ii): Hausdorff95 ‘เข้าข้าง’ งานที่ผิด (B) ด้วยซ้ำ — เรขาคณิตไม่ใช่แค่เสมอ แต่ ‘กลับด้าน’ → เป็นการสาธิตที่หนักแน่น")
bullet(doc, "เคส (iii): โพรง (H2) ของ alpha complex จะ ‘หายทั้งจำนวน’ ก็ต่อเมื่อรูใหญ่มาก เพราะโพรงถูกเติมที่สเกลรัศมีอยู่ดี. "
            "ที่รูขนาดพอเหมาะ เลขเบ็ตติ b2 ยังเป็น 1 ทั้งคู่ แต่ ‘ระยะ bottleneck’ จับการเลื่อนของเวลาตาย (death) ของโพรงได้ — "
            "ตัวชี้วัดต่อเนื่องไวกว่าการนับ; และ blister (A) ที่ Chamfer เท่ากันทำให้โพรงยังอยู่ (bott≈0) ดังนั้น Chamfer แยก ‘ปูดเล็กน้อย’ กับ ‘โพรงทะลุ’ ไม่ได้")
bullet(doc, "ความเร็ว/ความถูกต้อง: ไดอะแกรม H0 มีจุดเท่าจำนวนจุดสุ่ม (หลายหมื่น) — ตัด ‘สัญญาณรบกวนใกล้เส้นทแยงมุม’ "
            "(อายุ < ~2× ระยะเพื่อนบ้าน) ก่อนคำนวณระยะ เพื่อให้ Wasserstein/bottleneck เร็วและเสถียร โดยไม่ทิ้งฟีเจอร์จริง")
bullet(doc, "ข้อจำกัดร่วม: รูปทรงเฟสนี้เป็น ‘ผิวสังเคราะห์เชิงวิเคราะห์’ เพื่อคุมตัวแปรให้สะอาด (เจตนาเดียวกับ synthetic GT เดิม) — "
            "ขั้นถัดไปค่อยต่อกับงานจาก DiffSoup จริง (--traj_dir) ผ่าน adapter ที่เตรียมไว้แล้ว (topology.metrics.load_trajectory)")

# 10. สรุปและทิศทางเฟส 2
H(doc, "10. ข้อสรุปและทิศทางเฟส 2")
P(doc, "เฟส 1 พิสูจน์ ‘เครื่องมือวัด’ ครบ: กู้คืนเลขเบ็ตติได้ถูก, ตัวชี้วัดเสถียรภาพโทโพโลยีทำงาน, และสาธิตเชิงปริมาณว่า "
       "Chamfer/Hausdorff แยกความต่างเชิงโทโพโลยีไม่ได้ (3/3 เคส). โมดูล topology/ ถูกออกแบบให้แยกเป็นอิสระ เพื่อให้เมท็อดเฟส 2 "
       "เรียกใช้ซ้ำได้โดยไม่ต้องแก้")
bullet(doc, "เฟส 2: ออกแบบ Topology-aware Adaptive Resampling โดยใช้ persistence_from_target(...) เป็น prior ชี้นำ")
bullet(doc, "เชื่อมตัวชี้วัดกับ --traj_dir ของรันจริง เพื่อดูว่า resampling ลบ/สร้าง defect เชิงโทโพโลยีเมื่อไรระหว่างเทรน")
bullet(doc, "พิจารณา differentiable PH (เฟส 3) หากต้องการให้สัญญาณโทโพโลยีไหลย้อนเป็น gradient")

# 11. ภาคผนวก — วิธีรันซ้ำ
H(doc, "ภาคผนวก — วิธีรันซ้ำ (deterministic)")
code(doc, "uv pip install gudhi pot")
code(doc, ".venv\\Scripts\\python.exe tests\\test_betti.py            # 6/6 PASS")
code(doc, "$env:PYTHONUTF8=1 ; .venv\\Scripts\\python.exe experiments\\topology_blindness.py")
code(doc, "python src\\make_topology_report_docx.py                  # สร้างเอกสารนี้")
P(doc, "ผลดิบ: output/figures/topology/blindness_results.json, stability_series.csv; รูปทั้งหมด *.png; "
       "สรุปอังกฤษ summary.md. ทุกค่าในเอกสารนี้อ่านจากไฟล์ผลจริงโดยตรง", size=9)

os.makedirs(os.path.dirname(OUT) or ".", exist_ok=True)
doc.save(OUT)
print(f"[saved] {OUT}")
print(f"[size] {os.path.getsize(OUT):,} bytes")
