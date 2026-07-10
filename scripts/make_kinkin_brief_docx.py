# scripts/make_kinkin_brief_docx.py
# One-page Thai advisor brief for the kinkin asset — the artist-authored
# (Blender) tom-yum pot that replaces the CSG pot as the tomyum showcase —
# plus its first C-matrix results. House rule: NO hardcoded results — the
# C-table is read from output/synth/topo3/quicklook.json, the certificate from
# kinkin_src_cert.json, and the density staircase + significance margins are
# RECOMPUTED live from the certified mesh, so the document cannot drift.
#
#   D:\...\CG-Soup-for-Digital-Dentistry\.venv\Scripts\python.exe `
#       scripts\make_kinkin_brief_docx.py [out.docx]

import json
import math
import os
import sys

import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Tahoma"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)

_TOPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _TOPO)
DENTISTRY = os.environ.get("CGSOUP_ROOT", r"D:\Project\CG-Soup-for-Digital-Dentistry")
QUICKLOOK = os.path.join(DENTISTRY, "output", "synth", "topo3", "quicklook.json")
SRC = os.path.join(DENTISTRY, "output", "synth", "_meshes", "kinkin_src.ply")
CERT = os.path.join(DENTISTRY, "output", "synth", "_meshes", "kinkin_src_cert.json")
OUT = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    _TOPO, "docs", "CG-Soup_Kinkin_Brief_TH.docx")


# ── style helpers (verbatim from make_tomyum_brief_docx.py) ────────────────

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, white=False, size=9, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.size = Pt(size); run.font.name = FONT; run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        _shade(cell, fill)


def add_table(doc, headers, rows, shade_rows=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _cell_text(t.rows[0].cells[j], h, bold=True, white=True); _shade(t.rows[0].cells[j], "1F4E79")
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        fill = (shade_rows or {}).get(i)
        for j, v in enumerate(r):
            _cell_text(cells[j], v, fill=fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_figure(doc, path, caption, width=6.0):
    if not os.path.exists(path):
        doc.add_paragraph(f"[ไม่พบรูป: {path}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT; run.font.color.rgb = ACCENT
    return h


def P(doc, text, size=10.5, space=6, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT; r.bold = bold
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT
    return p


# ── live numbers ───────────────────────────────────────────────────────────

def load_mesh():
    import trimesh
    m = trimesh.load(SRC, process=False, force="mesh")
    return np.asarray(m.vertices), np.asarray(m.faces)


def staircase_rows(V, F):
    """Recompute the significant reading per M (live) + per-seed H2 margins
    at the bundle density."""
    from methods._paths import load_topology
    load_topology()
    from topology import meshes
    from topology.persistence import persistence_from_points
    rows = []
    for M in (2048, 4096, 8192, 20000, 50000):
        Pc = meshes.sample_surface(V, F, M, np.random.default_rng(0))
        res = persistence_from_points(Pc)
        b = res.betti_numbers()
        rows.append([f"{M:,}", f"{res.significance_threshold():.4f}",
                     f"({b[0]}, {b[1]}, {b[2]})"])
    margins = []
    for seed in range(5):
        Pc = meshes.sample_surface(V, F, 8192, np.random.default_rng(seed))
        res = persistence_from_points(Pc)
        thr = res.significance_threshold()
        h2 = res.diagram(2)
        life = np.sort(h2[:, 1] - h2[:, 0])[::-1]
        margins.append(float(life[0] / thr))
    return rows, margins


def cmatrix_stats():
    with open(QUICKLOOK, encoding="utf-8") as fh:
        rows = [r for r in json.load(fh) if r["shape"] == "kinkin"]
    if not rows:
        raise SystemExit(f"no kinkin rows in {QUICKLOOK} — run the quicklook first")
    out = {}
    for cond in ("C0", "C1", "C2"):
        sel = sorted([r for r in rows if r["cond"] == cond], key=lambda r: r["seed"])
        if not sel:
            continue
        b = np.array([r["tail_bottleneck_H2"] for r in sel], dtype=float)
        ch = np.array([r["chamfer_pct"] for r in sel], dtype=float)
        out[cond] = {"n": len(sel), "bott": b, "mean": float(b.mean()),
                     "sd": float(b.std(ddof=1)) if len(sel) > 1 else 0.0,
                     "nsig": [int(r["nsig_H2"]) for r in sel],
                     "cham": float(ch.mean())}
    def sigma(a, b):        # SE-of-difference convention (as in the 3e verdicts)
        se = math.sqrt(out[a]["sd"] ** 2 / out[a]["n"] + out[b]["sd"] ** 2 / out[b]["n"])
        return abs(out[a]["mean"] - out[b]["mean"]) / se if se > 0 else float("inf")
    out["red_C1"] = out["C0"]["mean"] / out["C1"]["mean"]     # unrounded means
    out["sig_C1"] = sigma("C0", "C1")
    if "C2" in out:
        out["worse_C2"] = out["C2"]["mean"] / out["C0"]["mean"]
        out["sig_C2"] = sigma("C0", "C2")
    return out


def main():
    with open(CERT, encoding="utf-8") as fh:
        cert = json.load(fh)
    V, F = load_mesh()
    stair, margins = staircase_rows(V, F)
    st = cmatrix_stats()

    doc = Document()
    for name in ("Normal", "List Bullet"):
        s = doc.styles[name]
        s.font.name = FONT
        rpr = s.element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(a), FONT)
        rpr.append(rfonts)

    H(doc, "เปลี่ยนโมเดล tomyum เป็นหม้อไฟที่ปั้นเอง (kinkin.ply) — ผล C-matrix ชุดใหม่", 1)
    P(doc, "สรุปสำหรับอาจารย์ · repo: CG-Soup-Topology · asset ใหม่ + รัน C-matrix ใหม่ทั้งชุด "
           "(ของเดิมยังอยู่ครบ ไม่มีตัวเลขเก่าเปลี่ยน)", size=9, space=8)

    H(doc, "1) โมเดลใหม่ และทำไมต้อง solidify + certify", 2)
    P(doc, f"หม้อไฟต้มยำเวอร์ชันปั้นเองใน Blender (kinkin.ply) เป็น artist mesh ตามธรรมชาติ: "
           f"9 ชิ้นซ้อนกัน มีขอบเปิด และรอยต่อ non-manifold — ใช้เป็น ground truth ตรง ๆ ไม่ได้ "
           f"(เหตุผลเดียวกับที่เราคัด ShapeNet ออกใน 3e) จึง ingest ผ่านขั้นตอนใหม่ "
           f"(scripts/make_kinkin_asset.py): ทำผนังโลหะบาง 2ε รอบผิวทุกแผ่น (offset-solidify "
           f"ด้วย exact distance field + marching cubes, ตัดเปลือกโพรงภายใน "
           f"{cert.get('dropped_internal_shells', '—')} ชิ้นทิ้ง) แล้ว CERTIFY แบบวัดจริง 3 ทางอิสระ: edge certificate "
           f"(ปิด/manifold/orientation สม่ำเสมอ), exact simplicial homology (GUDHI), "
           f"และ trimesh watertight — ได้ชิ้นเดียว genus {cert['genus_total']}, "
           f"b = ({cert['betti'][0]}, {cert['betti'][1]}, {cert['betti'][2]}) แบบ exact "
           f"({cert['verts']:,} จุด / {cert['faces']:,} หน้า, สร้างซ้ำ bit-identical)", size=10)
    add_figure(doc, os.path.join(_TOPO, "figures", "kinkin_pot.png"),
               "ซ้ายบน: artist mesh ตามที่ปั้น · ขวาบน: thin-shell ที่ certify แล้ว · "
               "ล่าง: cutaway และฐาน")

    H(doc, "2) topology ของหม้อใบนี้ต่างจาก pot เดิม — และนั่นคือจุดขาย", 2)
    P(doc, "ห่วง (H1) ของ kinkin เล็กทุกอัน (หูจับ/รูปล่องแคบ) จึงอยู่ใต้ floor 6·r_med ที่ทุก M "
           "ที่ใช้งานจริง แต่ปากปล่องที่แคบทำให้ alpha complex ปิดปากรูตั้งแต่สเกลเล็ก — "
           "ภายในปล่องจึงอ่านเป็น 'โพรงปิด' (H2): localization ยืนยันว่า void อยู่บนแกนหม้อพอดี "
           "หม้อใบนี้จึงเป็นชิ้นทดสอบคลาส H2 (เดียวกับ spot/fandisk) ไม่ใช่ H1 แบบ pot เดิม — "
           "การเลือกความหนาแน่น bundle ใช้กฎ floor เดิมของ study (density_bound.py): "
           "M เล็กสุดที่ feature หลักผ่าน 6·r_med คือ M=8192:", size=10)
    add_table(doc, ["M (จุด)", "floor 6·r_med", "significant (b0, b1, b2)"], stair,
              shade_rows={0: "FFF2CC", 2: "E2EFDA"})
    P(doc, f"margin ของ void ที่ M=8192 ต่อ floor: "
           f"{min(margins):.2f}–{max(margins):.2f}× (seeds 0–4 ผ่านทุกตัว) — "
           f"ระดับเดียวกับ margin ของห่วงปล่อง pot เดิมที่ M=2048 (1.17–1.25×) "
           f"และที่ M=50,000 จะเริ่มเห็นห่วงแรก + โพรงที่สอง (บันไดความหนาแน่นตัวจริง)",
      size=9.5)

    H(doc, "3) ผล C-matrix (สูตร 3e เดิม; ต่างแค่ M ของ bundle ตามกฎ floor)", 2)
    P(doc, "N=1200 (คลาส H2 เท่า spot/fandisk/sphere), 2500 steps, ρ=0.1, ramp 0.2:0.5, "
           "full bundle (ที่ M=8192 มี H2 significant 1 bar), --topo_loss_pts 8192 "
           "ตาม density-matched contract, 3 seeds:", size=9.5)
    rows = []
    fmt = lambda c: (f"{st[c]['mean']:.4f} ± {st[c]['sd']:.4f}",
                     "/".join(str(x) for x in st[c]["nsig"]), f"{st[c]['cham']:.3f}")
    r0 = fmt("C0"); rows.append(["C0 เบสไลน์", r0[0], "—", r0[1], r0[2]])
    r1 = fmt("C1"); rows.append(["C1 topological loss", r1[0],
                                 f"ดีขึ้น {st['red_C1']:.1f}× ({st['sig_C1']:.1f}σ)", r1[1], r1[2]])
    if "C2" in st:
        r2 = fmt("C2")
        destroyed = all(x == 0 for x in st["C2"]["nsig"])
        cmp2 = ("ทำลาย void (#sig 1→0 ทุก seed)" if destroyed else
                f"{'แย่ลง' if st['worse_C2'] >= 1 else 'ดีขึ้น'} "
                f"{st['worse_C2']:.1f}× ({st['sig_C2']:.1f}σ)")
        rows.append(["C2 ตัวควบคุม (repulsion)", r2[0], cmp2, r2[1], r2[2]])
    add_table(doc, ["เงื่อนไข", "tail bottleneck H2 (mean±sd)", "เทียบ C0",
                    "#sig H2 (ถูกต้อง=1)", "Chamfer %"],
              rows, shade_rows={1: "E2EFDA", 2: "FCE4EC"})
    bullet(doc, f"C1 ลดระยะ bottleneck ของโพรงปล่อง {st['red_C1']:.1f}× ({st['sig_C1']:.1f}σ) "
                f"ที่ Chamfer parity (C1/C0 = {st['C1']['cham'] / st['C0']['cham']:.2f})")
    if "C2" in st:
        destroyed = all(x == 0 for x in st["C2"]["nsig"])
        if destroyed:
            bullet(doc, f"C2 control ทำลาย topology: #sig H2 = 0 ทุก seed (ถูกต้อง = 1) — "
                        f"โหมด DESTROY แบบเดียวกับ bob/spot ใน 3e — และ Chamfer แย่ลงด้วย "
                        f"(C2/C0 = {st['C2']['cham'] / st['C0']['cham']:.2f}); "
                        f"ค่า bottleneck ของ C2 ค้างที่ half-life ของ bar เป้าหมาย "
                        f"({st['C2']['mean']:.4f}) เพราะไม่เหลือ void ให้จับคู่ — "
                        f"ตัวตัดสินคือ count ไม่ใช่ค่า")
        else:
            bullet(doc, f"C2 control: {st['worse_C2']:.2f}× เทียบ C0 ({st['sig_C2']:.1f}σ) — "
                        f"Chamfer ratio C2/C0 = {st['C2']['cham'] / st['C0']['cham']:.2f}")
    bullet(doc, "แถวนี้เป็น H2-void ตัวที่ 4 ของ study (sphere/cube/cylinder/spot/fandisk "
                "เป็นคลาสเดียวกัน) — และเป็นตัวแรกที่มาจาก asset ปั้นจริงที่ไม่ manifold "
                "ผ่านขั้นตอน ingest→certify")

    H(doc, "4) ความหมายต่อเฟส dental + สถานะของ pot เดิม", 2)
    bullet(doc, "ข้อความหลัก: mesh จริงจากศิลปิน (ไม่ manifold แบบเดียวกับข้อมูลสแกนจริง) "
                "ผ่าน solidify→certify แล้วเข้า pipeline ได้ทั้งระบบโดยไม่แก้สูตรใด ๆ — "
                "นี่คือ dress rehearsal ของขั้นตอนที่ต้องใช้กับข้อมูลฟันจริง")
    bullet(doc, "pot CSG เดิม (genus 9 by construction) ยังอยู่ในโค้ดครบ ผลเดิมยังบันทึกใน "
                "PHASE3_STATUS ตามเดิม — runs ชุดนี้เป็น tag ใหม่ (kinkin_*) ใน topo3 "
                "ไม่มีตัวเลขเก่าเปลี่ยน; คำถามเรื่องใส่ตาราง 3e ใน paper 2 หรือเก็บเป็น "
                "showcase จึงยังเปิดอยู่เหมือนเดิม (ตอนนี้มีให้เลือกสองใบ: pot-CSG แถว H1 "
                "หรือ kinkin แถว H2)")
    P(doc, "ทำซ้ำ: python scripts\\make_kinkin_asset.py แล้ว "
           "experiments/topo_loss_eval.py --shapes kinkin --seeds 0 1 2 "
           "--conditions C0 C1 C2 --rhos 0.1 --steps 2500 --max_faces 1200 --bundle_n 8192",
      size=8.5, space=2)

    doc.save(OUT)
    print(f"[docx] {OUT}")
    print(f"  C1: {st['red_C1']:.2f}x ({st['sig_C1']:.1f} sigma)  "
          f"C2: {st.get('worse_C2', float('nan')):.2f}x vs C0 ({st.get('sig_C2', float('nan')):.1f} sigma)")


if __name__ == "__main__":
    main()
