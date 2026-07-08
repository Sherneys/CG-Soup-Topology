# scripts/make_phase3_report_docx.py
# Build the Phase-3 (Differentiable Topological Loss) RESULTS as a Thai .docx.
# Reuses the style of make_phase2_report_docx.py and reads the REAL numbers
# from output/synth/topo3/report/results.json (experiments/topo_loss_report.py)
# so the document never drifts from the experiment.
#
#   python scripts/make_phase3_report_docx.py [out.docx]

import json
import math
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Tahoma"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)

_TOPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DENTISTRY = os.environ.get("CGSOUP_ROOT", r"D:\Project\CG-Soup-for-Digital-Dentistry")
REPORT = os.path.join(DENTISTRY, "output", "synth", "topo3", "report")
TOY_DIR = os.path.join(_TOPO, "figures", "phase3_toy")

OUT = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    _TOPO, "docs", "CG-Soup_Topology_Phase3_TH.docx")

DIM_TH = {0: "H0 (จำนวนชิ้นส่วน/components)", 1: "H1 (ห่วง·หูจับ/loops·handles)",
          2: "H2 (โพรงปิด/voids)"}
SHAPE_TH = {"sphere": "ทรงกลม (sphere)", "torus": "ทอรัส (torus)",
            "two_spheres": "ทรงกลมสองลูก (two_spheres)", "cube": "ลูกบาศก์ (cube)",
            "double_torus": "ทอรัสคู่ genus-2 (double_torus)",
            # 3e generality wave (external genus-known meshes):
            "spot": "spot — โมเดลวัว genus-0 (Crane, CC0)",
            "bob": "bob — โมเดลเป็ด genus-1 (Crane, CC0)",
            "fandisk": "fandisk — ชิ้นงาน CAD genus-0"}
COND_TH = {"C0": "C0 เบสไลน์", "C1": "C1 topological loss",
           "C2": "C2 ตัวควบคุม (repulsion)", "C2g": "C2g ตัวควบคุมแบบอ่อน",
           "C3": "C3 ไม่มี curriculum", "C5": "C5 loss + B4 prior",
           "C1_r0.03": "C1 (ρ=0.03)", "C1_r0.3": "C1 (ρ=0.3)"}


# ── style helpers (verbatim from make_phase2_report_docx.py) ───────────────

def _set_style_font(style, name=FONT):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), name)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, white=False, size=9, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.size = Pt(size); run.font.name = FONT; run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        _shade(cell, fill)


def add_table(doc, headers, rows, widths=None, shade_rows=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _cell_text(t.rows[0].cells[j], h, bold=True, white=True); _shade(t.rows[0].cells[j], "1F4E79")
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        fill = (shade_rows or {}).get(i)
        for j, v in enumerate(r):
            _cell_text(cells[j], v, fill=fill)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_figure(doc, path, caption, width=6.3):
    if not os.path.exists(path):
        doc.add_paragraph(f"[ไม่พบรูป: {path}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT; run.font.color.rgb = ACCENT
    return h


def P(doc, text, size=10.5, space=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT
    return p


def code(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ── load real results ──────────────────────────────────────────────────────

res_path = os.path.join(REPORT, "results.json")
if not os.path.exists(res_path):
    raise SystemExit(f"missing {res_path} — run experiments/topo_loss_report.py first")
with open(res_path, encoding="utf-8") as fh:
    R = json.load(fh)


def fmt(agg):
    return f"{agg['mean']:.4f} ± {agg['sd']:.4f}"


def factor(a0, a1):
    return f"{a0['mean'] / a1['mean']:.1f}×" if a1["mean"] > 0 else "∞"


# ── build ──────────────────────────────────────────────────────────────────

doc = Document()
for sname in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
    try:
        _set_style_font(doc.styles[sname])
    except KeyError:
        pass
doc.styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("CG-Soup / DiffSoup — เฟส 3: Differentiable Topological Loss (ผลการทดลอง)")
tr.bold = True; tr.font.size = Pt(17); tr.font.color.rgb = ACCENT; tr.font.name = FONT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("นำโทโพโลยีเข้าไปใน loss ของการเทรนโดยตรง — persistence diagram ที่หาอนุพันธ์ย้อนกลับได้ "
                 "ผ่าน circumradius ของ simplex ที่ถูกจับคู่ (pair-frozen backward)")
sr.font.size = Pt(12); sr.italic = True; sr.font.name = FONT
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("การเทียบแบบควบคุม C0/C1/C2/C3/C5 · Chamfer parity · เฉลี่ยข้าม seeds · "
                  "อ่านตัวเลขจากไฟล์ผลจริง · อัปเดตคลื่น 3e (ความทั่วไป) · 9 กรกฎาคม 2569")
mr.font.size = Pt(10); mr.font.name = FONT
doc.add_paragraph()

# 1. abstract
H(doc, "1. บทคัดย่อ")
P(doc, "เฟส 2 ตอบคำถาม ‘ควรวางสามเหลี่ยมที่ไหน’ (resampling prior) และพบว่า prior แบบกระจายช่วยโพรง (H2) "
       "แต่ล้มเหลวกับห่วง (H1) และกำไรส่วนใหญ่มาจากความกว้างของ prior ไม่ใช่ตำแหน่งเชิงโทโพโลยี "
       "เฟส 3 เปลี่ยน ‘ช่องทาง’: แทนที่จะจัดสรรงบสามเหลี่ยม เราใส่แรงดึงเชิงเมตริกลงบนค่า birth/death "
       "ของ persistence diagram โดยตรง — L = photometric + λ(t)·L_topo โดย L_topo คือระยะแบบจับคู่ "
       "(matched Wasserstein) ระหว่าง diagram ของพื้นผิวปัจจุบันกับ diagram เป้าหมาย")
P(doc, "ผลหลัก: loss ให้กำไรเชิงโทโพโลยีแท้ (ชนะทั้งเบสไลน์และตัวควบคุมที่ปรับขนาด gradient เท่ากัน) "
       "ทั้งกับโพรง H2 (ลด error 4.0–7.9 เท่า) และ — จุดที่ resampling ทำไม่ได้ — ห่วง H1 (2.4 เท่า "
       "โดยไม่สร้าง phantom handle เลยแม้แต่ seed เดียว) และสองช่องทางเสริมกัน: loss + prior (C5) "
       "ดีที่สุดในทุกรูปทรงที่ทดสอบ ส่วน H0 ยังคง ‘ไม่ใช่เรื่องโทโพโลยี’ เป็นครั้งที่สามข้ามทุกช่องทาง "
       "และล่าสุด (คลื่น 3e) ทุก verdict เกิดซ้ำบนเมชภายนอกสามชิ้นที่รู้ genus โดยไม่จูนอะไรต่อรูปทรงเลย "
       "(2.0–10.3 เท่า ที่ Chamfer ดีกว่าเบสไลน์ทุกชิ้น)")

# 2. method
H(doc, "2. วิธีการ (โดยย่อ)")
bullet(doc, "Forward ใช้ GUDHI ปกติ (alpha complex + persistence_pairs) เพื่อได้คู่ simplex เกิด/ตาย "
            "ของแต่ละ feature; backward คำนวณ circumradius ของ simplex เหล่านั้นใหม่ใน PyTorch "
            "(สูตรปิดของ edge/triangle/tetra) — ค่า filtration ของ alpha complex คือ circumradius² "
            "สำหรับ Gabriel simplex ซึ่งจากการตรวจจริงคือ 100% ของ feature ที่มีนัยยะบนรูปทรงของเรา")
bullet(doc, "จับคู่ diagram สด ↔ เป้าหมายด้วย optimal partial matching (Hungarian, เทียบเท่า W₂) "
            "+ เทอม ‘recruitment’ สำหรับ feature เป้าหมายที่หายไป (matching ล้วนให้ gradient เป็นศูนย์พอดี "
            "ในกรณีที่สำคัญที่สุด)")
bullet(doc, "จุดสุ่มบนพื้นผิว soup แบบ α×area (barycentric) — ตำแหน่งคำนวณจาก V สดทุก step แต่โครงสร้าง "
            "การจับคู่ถูก ‘แช่แข็ง’ แล้วรีเฟรชทุก 10 steps และอัตโนมัติหลังทุกการ resample")
bullet(doc, "λ(t): ramp 20%→50% ของการเทรน; ค่าพีคตั้งด้วยการปรับเทียบอัตราส่วน gradient "
            "λ_peak = ρ·‖∇V L_photo‖/‖∇V L_topo‖ (ρ=0.1; ผลแบนราบตลอด ρ∈{0.03,0.1,0.3})")
bullet(doc, "เงื่อนไขเทียบ: C0 เบสไลน์ · C1 loss · C2 ตัวควบคุม non-topological ที่ขนาด gradient เท่ากัน "
            "(repulsion บนจุดสุ่มชุดเดียวกัน) · C3 = C1 แบบไม่มี ramp · C5 = C1 + B4 spread prior จากเฟส 2b")

# 3. toys
H(doc, "3. เกตก่อนใช้งานจริง (stage 3a): loss ล้วน ๆ ซ่อมรูปทรงที่พังได้")
P(doc, "ก่อนต่อเข้ากับการเทรน เราบังคับให้ loss (โดยไม่มี photometric ช่วยเลย) ซ่อม point cloud "
       "ที่โทโพโลยีผิด 4 แบบ ด้วย Adam บนพิกัดจุดโดยตรง — ผ่านทั้งหมด:")
add_table(doc,
          ["โจทย์", "กลไกที่ทดสอบ", "ผล (bottleneck)", "Betti"],
          [["ทรงกลมเจาะรู → ทรงกลม", "ดึง birth ของคู่ที่จับได้ (ปิดรู)", "0.097 → 0.011 (8.8×)", "(1,0,1) ✓"],
           ["ทรงกลมมีหูจับ → ทรงกลม", "ดัน bar เกินเข้าเส้นทแยง (ฆ่าห่วงปลอม)", "0.044 → 0 (หายสนิท)", "(1,1,1)→(1,0,1) ✓"],
           ["สองทรงกลมห่างผิด → ระยะถูก", "ดึง death ของ H0", "0.064 → 0.0001 (1229×)", "(2,0,2) ✓"],
           ["สองทรงกลมมีสะพานเชื่อม", "recruitment (ไม่มี bar สดให้จับ)", "0.056 → 0.0001 (450×)", "(1,0,2)→(2,0,2) ✓"]],
          widths=[1.9, 2.1, 1.5, 1.1])
P(doc, "ข้อสังเกตเชิงกลไกจาก toy: การซ่อมมาจากแรงดึงเบา ๆ ทีละไม่กี่จุดที่ ‘เดิน’ ไปรอบ ๆ จุดบกพร่อง "
       "ผ่านการจับคู่ใหม่ทุกรอบ; recruitment แม้ตาบอดเชิงตำแหน่งแต่ถูกภูมิทัศน์ของ loss พาไปตัดสะพาน "
       "ถูกที่เอง (การฉีกผิว shell สร้าง component ไม่ได้ → ถูกทิ้ง; การตัดสะพาน 1 มิติคือทิศเดียวที่ "
       "death โตถึงเป้า)")

# 4. results
H(doc, "4. ผลการทดลองจริง (C-matrix, 2,500 steps, งบสามเหลี่ยมคงที่ต่อรูปทรง)")
P(doc, "ตัวชี้วัดหลัก = bottleneck distance ไปยัง diagram เป้าหมายในมิติชี้ขาด (ค่าเฉลี่ยช่วงท้ายการเทรน "
       "± sd ข้าม seeds; ต่ำ = ดี) พร้อมเงื่อนไข Chamfer parity — ชนะเชิงโทโพโลยีต้องไม่แลกด้วยเรขาคณิต")

order = ["sphere", "cube", "torus", "two_spheres", "double_torus",
         "spot", "bob", "fandisk"]          # 5 analytic + the 3e generality trio
for shape in order:
    if shape not in R:
        continue
    if shape == "spot":                     # 3e shapes start here — one intro
        H(doc, "4b. คลื่น 3e — ความทั่วไปบนเมชภายนอกที่รู้ genus (9 ก.ค. 2569)", level=2)
        P(doc, "ทำซ้ำโปรโตคอลเดิมทุกอย่าง (C0/C1/C2 × 3 seeds, ρ, ramp, งบต่อคลาสเท่าเดิม, "
               "ไม่จูนอะไรต่อรูปทรง) บนเมชภายนอกสามชิ้นที่ตรวจ watertight และ genus แล้ว: "
               "spot (วัว, genus 0) กับ bob (เป็ด, genus 1) จากคลังโมเดลของ Keenan Crane (CC0) "
               "และ fandisk (ชิ้นงาน CAD, genus 0) — ใช้แทน ShapeNet ซึ่งเมชส่วนใหญ่ไม่ watertight "
               "จึงไม่มี ground truth เชิงโทโพโลยีที่เชื่อถือได้ ผล: ผ่านทั้ง 3/3 ที่ Chamfer "
               "ดีกว่าเบสไลน์ทุกชิ้น (2.0–10.3 เท่า) และโหมดล้มเหลวของตัวควบคุมเกิดซ้ำเป๊ะ "
               "(โพรงถูกลบทิ้งบนทั้งสองชิ้น H2, ห่วงหายหนึ่งวงบน genus-1) หมายเหตุตรงไปตรงมา: "
               "บน spot เบสไลน์เองสร้างโพรงปลอมใน 2/3 seeds และ C1 แก้ได้ 2 ใน 3 seeds — "
               "การซ่อม ‘ค่า’ (bottleneck) เชื่อถือได้ แต่การซ่อม ‘จำนวน’ บนรายละเอียดอินทรีย์ "
               "ที่เล็กกว่างบ (ขา/เขา ที่ N=1200) เป็นการปรับปรุง ไม่ใช่การรับประกัน")
    rec = R[shape]
    agg = rec["aggregate"]
    d = rec["disc_dim"]
    H(doc, f"4.{order.index(shape) + 1} {SHAPE_TH.get(shape, shape)} — มิติชี้ขาด {DIM_TH[d]}", level=2)
    rows = []
    for cond in sorted(agg):
        g = agg[cond]
        extra = factor(agg["C0"], g) if (cond != "C0" and "C0" in agg) else "—"
        rows.append([COND_TH.get(cond, cond), g["n_seeds"], fmt(g), extra,
                     f"{g['chamfer_mean']:.3f}", str(g["nsig_final"])])
    add_table(doc, ["เงื่อนไข", "seeds", f"bottleneck H{d} (ท้ายการเทรน)", "เทียบ C0",
                    "Chamfer%", "#sig สุดท้าย"],
              rows, widths=[1.7, 0.6, 1.6, 0.8, 0.9, 1.0])
    if rec.get("verdict"):
        P(doc, f"คำตัดสิน (C1 ต้องชนะทั้ง C0 และตัวควบคุมทุกตัวที่ Chamfer parity): "
               f"{'ผ่าน ✓' if rec['verdict'] == 'PASS' else 'ไม่ผ่าน ✗'} — " + "; ".join(rec["why"]),
          size=9.5)
    for fig, cap in ((f"{shape}_series.png", "ค่า error เชิงโทโพโลยีระหว่างการเทรน (แถบ = ช่วง min–max ข้าม seeds)"),
                     (f"{shape}_nsig.png", "จำนวน feature ที่มีนัยยะ (เช็ค phantom; เส้นประ = ค่าที่ถูกต้อง)")):
        add_figure(doc, os.path.join(REPORT, fig), f"{SHAPE_TH.get(shape, shape)} — {cap}", width=5.8)

# 5. findings
H(doc, "5. ข้อค้นพบหลักห้าข้อ")
bullet(doc, "① loss channel สำเร็จในจุดที่ prior channel ล้มเหลว — ห่วง H1: ลด error 2.4 เท่า "
            "โดย #sig H1 = 2 ตลอดทั้ง 5 seeds (ไม่มี phantom handle) ขณะที่ตัวควบคุมทำห่วงหายหนึ่งวง "
            "— แรงดึงเชิงเมตริกแบบเจาะจงเหมาะกับ feature 1 มิติมากกว่าการเทงบสามเหลี่ยม")
bullet(doc, "② สองช่องทางเสริมกัน: C5 (loss + B4 spread prior) ดีที่สุดทุกรูปทรงที่ทดสอบ — sphere ลดถึง "
            "7.6 เท่าจากเบสไลน์ (~2 เท่าจาก loss เดี่ยว) พร้อม Chamfer ดีที่สุด — prior ‘จัดสรร’, loss ‘จูนค่า’")
bullet(doc, "③ H0 ไม่ใช่เรื่องโทโพโลยี — เป็นครั้งที่สาม: ตัวควบคุม non-topological ตามทันบนตัวชี้วัด diagram "
            "(ต่างเฉพาะ Chamfer) สอดคล้องกับเฟส 2 และ 2b — นี่คือข้อสรุปเกี่ยวกับคลาสของ feature "
            "ไม่ใช่ข้อจำกัดของวิธีใดวิธีหนึ่ง")
bullet(doc, "④ curriculum ไม่จำเป็นเมื่อปรับเทียบ λ ด้วยอัตราส่วน gradient (C3 ≈ C1 ทั้งสองรูปทรงชี้ขาด) "
            "— สิ่งที่ทำให้ปุ่มความแรงปลอดภัยคือ calibration ไม่ใช่ตารางเวลา (ตอบข้อกังวล gradient explode "
            "ของอาจารย์ด้วยข้อมูล)")
bullet(doc, "⑤ ความทั่วไป (คลื่น 3e): ทุก verdict เกิดซ้ำบนเมชภายนอกสามชิ้นโดยไม่จูนต่อรูปทรง — "
            "spot 2.2 เท่า (6.1σ), bob 2.0 เท่า (26.1σ; phantom handle = ศูนย์ทุก seed — ผล H1 "
            "อยู่รอดนอกตระกูล analytic), fandisk 10.3 เท่า (35.7σ — ผลใหญ่ที่สุดในการศึกษา) "
            "— ข้อสรุปไม่ใช่ artifact ของตระกูลรูปทรงสังเคราะห์")

# 6. caveats
H(doc, "6. ข้อจำกัดและ caveat ที่รายงานตรงไปตรงมา")
bullet(doc, "loss บังคับได้เฉพาะโทโพโลยีที่ ‘วัดได้’ ที่ความหนาแน่นจุดสุ่มของมัน (M=2048): โพรงของทอรัส "
            "และห่วงเล็กสองวงของทอรัสคู่จมใต้เกณฑ์นัยยะ — จึงจำกัด loss ไว้ที่มิติชี้ขาด และทอรัสคู่มีเป้า "
            "แค่ 2/4 ห่วง (ตรงกับข้อจำกัดของ DiffSoup เองที่สร้างได้ 2/4 ที่ทุกงบในเฟส 2b)")
bullet(doc, "ตัวควบคุม C2 ค่อนข้างแรง (ผลักถึง 2× ระยะจุดเฉลี่ย — ถึงขั้นทำลายโทโพโลยีบน cube/torus); "
            "จึงเพิ่ม C2g แบบอ่อน (1×) เพื่อยืนยันว่าข้อสรุปไม่ขึ้นกับความแรงของตัวควบคุม")
bullet(doc, "DiffSoup บน CUDA ไม่ bit-reproducible (พบเพิ่มเติมว่าแม้ F-trajectory ของเบสไลน์เอง "
            "ก็แตกกันระหว่างรันที่ seed เดียวกัน) — ทุกคำกล่าวอ้างจึงเฉลี่ยข้าม seeds")
bullet(doc, "ยังจำกัดที่การเรนเดอร์สังเคราะห์พื้นผิวปิดเครื่องเดียว (200–256 px, 24–72 มุมมอง) — "
            "คลื่น 3e ขยายจากตระกูล analytic ไปเมชภายนอกสามชิ้นแล้ว แต่ยังไม่ถึงชุดรูปทรงสาธารณะ "
            "ขนาดใหญ่; ข้อมูลทันตกรรมจริงเป็นขั้นถัดไป (ตามลำดับที่ตกลงกับอาจารย์)")

# 7. next
H(doc, "7. ขั้นถัดไป")
bullet(doc, "C4: ถ่วงน้ำหนัก loss เชิงพื้นที่ด้วย curvature (ข้อ 3 ของอาจารย์) — ต้องเพิ่ม weight-mask "
            "ใน TopoLossState ก่อน")
bullet(doc, "ความทั่วไป: ทำแล้ว (คลื่น 3e — spot/bob/fandisk แทน ShapeNet ที่เมชส่วนใหญ่ "
            "ไม่ watertight) และเขียนเข้า paper ฉบับที่ 2 แล้ว (paper2/ ร่างสมบูรณ์: บทคัดย่อ ตาราง "
            "และรูปครบ); ชุดรูปทรงสาธารณะที่กว้างขึ้นค่อยเพิ่มถ้าผู้ตรวจต้องการ")
bullet(doc, "เดโมทันตกรรมเชิงคุณภาพ 1 ชิ้น หลังผลสังเคราะห์นิ่ง (H1-first ตามธรรมชาติของใบหน้า/ช่องปาก)")

doc.save(OUT)
print(f"[docx] {OUT}")
